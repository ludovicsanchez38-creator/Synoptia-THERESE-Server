"""
THERESE v2 - Word Document Generator Skill

Génère des documents Word (.docx) avec le style Synoptia.
Approche code-execution avec fallback parser legacy.
"""

import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.services.skills.base import FileFormat, SkillParams, SkillResult
from app.services.skills.code_executor import CodeGenSkill

logger = logging.getLogger(__name__)


# Palette Synoptia
SYNOPTIA_COLORS = {
    "background": RGBColor(0x0B, 0x12, 0x26),
    "text": RGBColor(0xE6, 0xED, 0xF7),
    "muted": RGBColor(0xA9, 0xB8, 0xD8),
    "primary": RGBColor(0x17, 0x33, 0xA6),
    "accent_cyan": RGBColor(0x22, 0xD3, 0xEE),
    "accent_magenta": RGBColor(0xE1, 0x1D, 0x8D),
    # Couleurs claires pour documents imprimables
    "heading": RGBColor(0x0F, 0x1E, 0x6D),
    "body": RGBColor(0x1A, 0x1A, 0x2E),
}


class DocxSkill(CodeGenSkill):
    """
    Skill de génération de documents Word.

    Crée des documents .docx professionnels avec le style Synoptia.
    Approche code-execution : le LLM génère du code python-docx.
    Fallback automatique vers l'ancien parser Markdown.
    """

    skill_id = "docx-pro"
    name = "Document Word Professionnel"
    description = "Génère un document Word structuré avec le style Synoptia"
    output_format = FileFormat.DOCX

    def __init__(self, output_dir: Path):
        super().__init__(output_dir)

    def get_system_prompt_addition(self) -> str:
        """Instructions pour le LLM : générer du code Python python-docx."""
        return """
## Instructions pour génération de document Word

Tu dois générer un **bloc de code Python** complet utilisant la bibliothèque `python-docx`.
Le code sera exécuté dans un environnement sandboxé avec les variables suivantes pré-injectées :
- `output_path` (str) : chemin où sauvegarder le fichier .docx
- `title` (str) : sujet demandé par l'utilisateur (utilise-le comme inspiration, mais génère un titre professionnel et pertinent pour le heading principal, PAS le prompt brut)
- `SYNOPTIA_COLORS` (dict) : palette de couleurs Synoptia

### Imports disponibles
docx (Document), docx.shared (Cm, Pt, Inches, RGBColor), docx.enum.text (WD_ALIGN_PARAGRAPH), docx.enum.style (WD_STYLE_TYPE), docx.enum.table, docx.oxml.ns (qn), datetime, json, re, math, Decimal

### Règles impératives
1. **Police** : Calibri 11pt pour le corps, Outfit pour les titres
2. **Titres** : `doc.add_heading(text, level=0/1/2/3)` avec couleurs Synoptia
   - Heading 0 : titre principal, 28pt, bold, couleur #0F1E6D
   - Heading 1 : 24pt, bold, couleur #0F1E6D
   - Heading 2 : 18pt, bold, couleur #1733A6
   - Heading 3 : 14pt, bold, couleur #0F1E6D
3. **Mise en forme** :
   - **Gras** : `run.bold = True`
   - *Italique* : `run.italic = True`
   - Listes à puces : `doc.add_paragraph(text, style='List Bullet')`
   - Listes numérotées : `doc.add_paragraph(text, style='List Number')`
4. **Tableaux** : `doc.add_table(rows, cols, style='Table Grid')` avec style professionnel
   - Header row en bold avec fond coloré
   - Bordures fines
5. **Mise en page** :
   - Marges : 2.5 cm (ou Inches(1))
   - Espacement : line_spacing = 1.15, space_after = Pt(10)
6. **Footer** : "Généré par THÉRÈSE - Synoptïa" centré, 9pt, italique, couleur #A9B8D8
7. **Structure** :
   - Introduction (contexte, objectif)
   - Corps (sections logiques avec titres)
   - Conclusion (résumé, prochaines étapes)
8. **Finir par** : `doc.save(output_path)`

IMPORTANT : Ne pas ajouter de bloc récapitulatif (Sujet / Action / Date) à la fin du document. Le document se termine avec le contenu demandé uniquement (ou le footer Synoptïa si applicable). Ce bloc est réservé aux réponses chat.

### Structure du code

```python
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Configuration styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Titre principal
doc.add_heading(title, level=0)

# ... contenu du document ...

# Footer
section = doc.sections[-1]
footer = section.footer
footer.is_linked_to_previous = False
para = footer.paragraphs[0]
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run("Généré par THÉRÈSE - Synoptïa")
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0xA9, 0xB8, 0xD8)

doc.save(output_path)
```

Génère UNIQUEMENT le bloc ```python``` avec le code complet. Pas d'explication avant ou après.
"""

    def get_markdown_prompt_addition(self) -> str:
        """Instructions Markdown pour modèles non code-capable."""
        return """
## Instructions pour génération de document Word

Génère le contenu du document en Markdown bien structuré.
Utilise :
- # Titre principal (un seul, en début de document)
- ## Sections principales
- ### Sous-sections
- Listes à puces (- item) et listes numérotées (1. item)
- **gras** pour les termes importants
- *italique* pour les nuances
- Tableaux Markdown (| col1 | col2 |) pour les données tabulaires

### Structure attendue
1. Introduction (contexte, objectif)
2. Corps (sections logiques avec titres)
3. Conclusion (résumé, prochaines étapes)

NE génère PAS de code Python. Écris directement le contenu textuel du document.
"""

    async def _fallback_execute(
        self, params: SkillParams, file_id: str, output_path: Path
    ) -> SkillResult:
        """
        Fallback : ancien parser Markdown -> python-docx.

        Args:
            params: Paramètres de génération
            file_id: ID du fichier pré-généré
            output_path: Chemin de sortie pré-calculé

        Returns:
            Résultat avec chemin vers le fichier généré
        """
        # Créer le document
        doc = Document()

        # Appliquer les styles Synoptia
        self._setup_styles(doc)

        # Ajouter le titre principal
        title_para = doc.add_heading(params.title, level=0)
        self._style_title(title_para)

        # Parser et ajouter le contenu
        self._add_content(doc, params.content)

        # Ajouter le footer Synoptia
        self._add_footer(doc)

        # Sauvegarder
        doc.save(str(output_path))

        # Calculer la taille
        file_size = output_path.stat().st_size

        logger.info(f"Generated DOCX (fallback): {output_path} ({file_size} bytes)")

        return SkillResult(
            file_id=file_id,
            file_path=output_path,
            file_name=output_path.name,
            file_size=file_size,
            mime_type=self.get_mime_type(),
            format=self.output_format,
        )

    def _setup_styles(self, doc: Document) -> None:
        """Configure les styles du document."""
        styles = doc.styles

        # Style Normal
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.font.color.rgb = SYNOPTIA_COLORS["body"]
        normal.paragraph_format.space_after = Pt(10)
        normal.paragraph_format.line_spacing = 1.15

        # Style Heading 1
        if "Heading 1" in styles:
            h1 = styles["Heading 1"]
            h1.font.name = "Outfit"
            h1.font.size = Pt(24)
            h1.font.bold = True
            h1.font.color.rgb = SYNOPTIA_COLORS["heading"]
            h1.paragraph_format.space_before = Pt(24)
            h1.paragraph_format.space_after = Pt(12)

        # Style Heading 2
        if "Heading 2" in styles:
            h2 = styles["Heading 2"]
            h2.font.name = "Outfit"
            h2.font.size = Pt(18)
            h2.font.bold = True
            h2.font.color.rgb = SYNOPTIA_COLORS["primary"]
            h2.paragraph_format.space_before = Pt(18)
            h2.paragraph_format.space_after = Pt(8)

        # Style Heading 3
        if "Heading 3" in styles:
            h3 = styles["Heading 3"]
            h3.font.name = "Outfit"
            h3.font.size = Pt(14)
            h3.font.bold = True
            h3.font.color.rgb = SYNOPTIA_COLORS["heading"]
            h3.paragraph_format.space_before = Pt(14)
            h3.paragraph_format.space_after = Pt(6)

    def _style_title(self, para) -> None:
        """Applique le style au titre principal."""
        for run in para.runs:
            run.font.name = "Outfit"
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = SYNOPTIA_COLORS["heading"]

    def _add_content(self, doc: Document, content: str) -> None:
        """
        Parse le contenu Markdown et l'ajoute au document.

        Args:
            doc: Document Word
            content: Contenu en format Markdown
        """
        # Supprimer les blocs de code résiduels (```python...```)
        content = re.sub(
            r"```(?:python|py|javascript|js|bash|sh|json|xml|html|css|sql|yaml|yml)?\s*\n.*?(?:```|$)",
            "",
            content,
            flags=re.DOTALL,
        )

        lines = content.split('\n')
        list_number = 0
        in_code_block = False

        for line in lines:
            line = line.strip()
            if not line:
                list_number = 0
                continue

            # Ignorer les lignes résiduelles de clôture de code
            if line.startswith('```'):
                in_code_block = not in_code_block
                continue
            if in_code_block:
                continue

            # Headings
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            # Listes à puces
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:]
                para = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(para, text)
            # Listes numérotées
            elif re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line)
                para = doc.add_paragraph(style='List Number')
                self._add_formatted_text(para, text)
                list_number += 1
            # Tableaux (simplifié)
            elif line.startswith('|'):
                # Ignorer les lignes de séparation
                if not re.match(r'^\|[\s\-:|]+\|$', line):
                    cells = [c.strip() for c in line.split('|')[1:-1]]
                    if cells:
                        para = doc.add_paragraph()
                        para.add_run('\t'.join(cells))
            # Paragraphe normal
            else:
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)

    def _add_formatted_text(self, para, text: str) -> None:
        """
        Ajoute du texte formaté (gras, italique) à un paragraphe.

        Args:
            para: Paragraphe Word
            text: Texte avec formatage Markdown
        """
        # Pattern pour **gras**, *italique*, `code`
        pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`|[^*`]+)'
        parts = re.findall(pattern, text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = para.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = para.add_run(part[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            else:
                para.add_run(part)

    def _add_footer(self, doc: Document) -> None:
        """Ajoute un footer Synoptia au document."""
        section = doc.sections[-1]
        footer = section.footer
        footer.is_linked_to_previous = False

        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run("Généré par THÉRÈSE - Synoptïa")
        run.font.size = Pt(9)
        run.font.color.rgb = SYNOPTIA_COLORS["muted"]
        run.font.italic = True
